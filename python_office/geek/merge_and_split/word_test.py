'''
直接安装不成功，可以先下载再安装 安装包地址 https://pypi.org/project/python-docx/#files
word 文件操作的模块名称叫  python-doxc
该模块安装依赖 lxml 如果网络问题安装失败，可用国内镜像加速，如：
pip install lxml -i http://pypi.douban.com/simple/ --trusted-host pypi.douban.com
然后用下载的离线包安装docx
pip install C:\\Users\\admin\\Downloads\\python-docx-0.8.11.tar.gz
'''
import docx
from docx import Document
from pathlib import Path, PurePath

# word文件所在路径
word_files_path = 'word'

# 取得该目录下所有的docx格式文件
p = Path(word_files_path)
# 对文件以后缀进行过滤
files = [x for x in p.iterdir() if PurePath(x).match('*.docx')]

'''
Python 3.5 新增类型提示功能
    冒号后面是建议传入的参数类型
    箭头后面是建议函数返回的类型
'''
# 只获取内容进行合并
def merge_without_format(docx_files: list):
    # 实例化文档对象
    doc = Document()

    # 遍历每个文件 这里调用了排序函数排序
    for docx_file in sorted(docx_files):
        another_doc = Document(docx_file)
        # 获取每个文件的所有“段落”
        paras = another_doc.paragraphs
        # 获取所有段落的文字内容
        # paras_content = [para.text for para in paras]
        print('文件 %s 的文件内容' % docx_file)
        for para in paras:
            # 为新的word文件创建一个新段落
            newpar = doc.add_paragraph('')
            # 将提取的内容写入新的文本段落中
            text = para.text
            newpar.add_run(text)
            print(text)

    # 所有文件合并完成后在指定路径进行保存
    doc.save(Path(word_files_path, '合并word文字内容.docx'))

# 完全合并，不管内容是文字、图片还是其它
def merge_files(docx_files:list):
    doc = Document()
    for docx_file in sorted(docx_files):
        another_doc = Document(docx_file)
        # 遍历每一页
        for word_body in another_doc.element.body:
            # 合并内容到新的word文档
            doc.element.body.append(word_body)

    doc.save(Path(word_files_path,'合并word所有内容.docx'))

# 调用函数
# merge_without_format(files)
# merge_files(files)


from docx.shared import RGBColor

'''
该方法主要将txt等其它文件中的文本内容输出到word格式文件中，并设置相应的字体格式等
属性设置参考：https://python-docx.readthedocs.io/en/latest/api/text.html#run-objects
'''
def add_content_from_txt(txt_path: str):
    # 实例化文档对象
    doc = Document()
    # 像文本对象中添加内容
    file = open(txt_path, encoding='utf-8')
    content = file.read()
    para = doc.add_paragraph().add_run(content)
    # 字体设置
    para.font.name = '仿宋'
    para.font.underline = True
    para.font.color.rgb = RGBColor(255, 128, 128)

    doc.save(Path(word_files_path,'合并txt并格式化的word文件.docx'))

# add_content_from_txt('word/test.txt')


from docx import shared

def add_picture(img_path:str):
    doc = Document()
    # 添加图片函数，如果没有指定width和height，会将原图添加到文档中。有指定则会按比例缩放。
    doc.add_picture(img_path, width=shared.Inches(5),height=shared.Pt(200))
    doc.save(Path(word_files_path,'图片合并到word中测试文件.docx'))


add_picture('word/img.png')